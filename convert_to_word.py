from docx import Document
from docx.shared import Pt
import os

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # 設定字體樣式 (預設支持微軟正黑體等中文字體)
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(12)

    if not os.path.exists(md_path):
        print(f"找不到檔案: {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 處理標題
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        # 處理清單 (簡單處理)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # 處理普通段落
        else:
            p = doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"成功將 {md_path} 轉換為 {docx_path}")

if __name__ == "__main__":
    markdown_to_docx('Learning.md', 'Learning_Note.docx')
